#!/usr/bin/env python3
"""
Resume Updater Script
Intelligently updates your DevOps resume based on job requirements
Usage: python3 resume_updater.py
"""

from docx import Document
from docx.shared import Pt
import os
import sys
from datetime import datetime

class ResumeUpdater:
    def __init__(self, original_resume_path, output_dir):
        self.original_resume_path = original_resume_path
        self.output_dir = output_dir
        self.doc = None
        
        # Common DevOps technologies that should be bold
        self.tech_terms = [
            'AWS', 'Azure', 'GCP', 'Google Cloud',
            'ECS', 'Fargate', 'Lambda', 'EC2', 'S3', 'VPC', 'ELB', 'CloudFormation',
            'Aurora', 'PostgreSQL', 'MySQL', 'DynamoDB', 'MongoDB', 'RDS',
            'Kubernetes', 'Docker', 'OpenShift', 'Helm', 'ArgoCD', 'Kustomize',
            'Jenkins', 'GitLab CI/CD', 'GitHub Actions', 'Tekton', 'Bamboo', 'TeamCity',
            'Terraform', 'Pulumi', 'Ansible', 'Chef', 'Puppet',
            'Prometheus', 'Grafana', 'DataDog', 'Splunk', 'ELK', 'Nagios',
            'Python', 'Bash', 'Shell', 'Groovy', 'Go', 'Ruby', 'Perl',
            'Apache Kafka', 'Kinesis', 'RabbitMQ', 'Redis',
            'Nginx', 'Apache', 'Tomcat', 'JBoss', 'WebSphere',
            'Git', 'GitHub', 'GitLab', 'Bitbucket', 'SVN',
            'JIRA', 'Confluence', 'ServiceNow',
            'Linux', 'RHEL', 'CentOS', 'Ubuntu', 'Debian',
            'Maven', 'Ant', 'Gradle', 'npm', 'yarn',
            'Istio', 'Linkerd', 'Flux', 'Crossplane',
            'SonarQube', 'Trivy', 'Snyk', 'Checkmarx',
            'CI/CD', 'DevOps', 'GitOps', 'DevSecOps', 'SRE',
            'REST API', 'GraphQL', 'gRPC', 'WebSockets',
            'Microservices', 'BFF', 'API Gateway',
            'SAFe Agile', 'Scrum', 'Kanban',
            'SSL', 'TLS', 'mTLS', 'OAuth', 'SAML',
            'CloudWatch', 'Application Insights', 'New Relic',
            'Selenium', 'JUnit', 'TestNG', 'pytest',
            'YAML', 'JSON', 'XML', 'HCL'
        ]
    
    def load_resume(self):
        """Load the original resume"""
        if not os.path.exists(self.original_resume_path):
            print(f'Error: Resume file not found at {self.original_resume_path}')
            sys.exit(1)
        
        self.doc = Document(self.original_resume_path)
        print(f'✓ Loaded resume: {os.path.basename(self.original_resume_path)}')
    
    def parse_requirements(self, requirements_text):
        """Parse job requirements and extract key technologies and skills"""
        print('\n📋 Analyzing requirements...')
        
        requirements = {
            'cloud_services': [],
            'containers': [],
            'cicd_tools': [],
            'programming': [],
            'databases': [],
            'monitoring': [],
            'messaging': [],
            'other_skills': [],
            'methodologies': []
        }
        
        req_lower = requirements_text.lower()
        
        # Cloud services
        if 'ecs' in req_lower or 'fargate' in req_lower:
            requirements['cloud_services'].append('ECS Fargate')
        if 'lambda' in req_lower or 'serverless' in req_lower:
            requirements['cloud_services'].append('Lambda')
        if 'aurora' in req_lower:
            requirements['cloud_services'].append('Aurora PostgreSQL')
        if 'dynamodb' in req_lower:
            requirements['cloud_services'].append('DynamoDB')
        if 'kinesis' in req_lower:
            requirements['cloud_services'].append('Kinesis')
        if 'codepipeline' in req_lower or 'aws pipeline' in req_lower:
            requirements['cloud_services'].append('AWS CodePipeline')
        if 'api gateway' in req_lower:
            requirements['cloud_services'].append('API Gateway')
        
        # Containers & Orchestration
        if 'kubernetes' in req_lower or 'k8s' in req_lower:
            requirements['containers'].append('Kubernetes')
        if 'docker' in req_lower:
            requirements['containers'].append('Docker')
        if 'helm' in req_lower:
            requirements['containers'].append('Helm')
        if 'argocd' in req_lower:
            requirements['containers'].append('ArgoCD')
        
        # CI/CD
        if 'tekton' in req_lower:
            requirements['cicd_tools'].append('Tekton')
        if 'github actions' in req_lower:
            requirements['cicd_tools'].append('GitHub Actions')
        if 'gitlab' in req_lower:
            requirements['cicd_tools'].append('GitLab CI/CD')
        if 'jenkins' in req_lower:
            requirements['cicd_tools'].append('Jenkins')
        
        # Databases
        if 'postgres' in req_lower:
            requirements['databases'].append('PostgreSQL')
        if 'mysql' in req_lower:
            requirements['databases'].append('MySQL')
        if 'mongodb' in req_lower:
            requirements['databases'].append('MongoDB')
        
        # Messaging
        if 'kafka' in req_lower:
            requirements['messaging'].append('Apache Kafka')
        if 'kinesis' in req_lower and 'Kinesis' not in requirements['cloud_services']:
            requirements['messaging'].append('AWS Kinesis')
        
        # Monitoring
        if 'prometheus' in req_lower:
            requirements['monitoring'].append('Prometheus')
        if 'grafana' in req_lower:
            requirements['monitoring'].append('Grafana')
        if 'datadog' in req_lower:
            requirements['monitoring'].append('DataDog')
        
        # Architecture patterns
        if 'microservices' in req_lower or 'micro services' in req_lower or 'micro-services' in req_lower:
            requirements['other_skills'].append('microservices')
        if 'bff' in req_lower or 'backend for frontend' in req_lower:
            requirements['other_skills'].append('BFF')
        if 'event-driven' in req_lower or 'event driven' in req_lower:
            requirements['other_skills'].append('event-driven architecture')
        if 'service mesh' in req_lower:
            requirements['other_skills'].append('service mesh')
        
        # Methodologies
        if 'safe' in req_lower or 'safe agile' in req_lower:
            requirements['methodologies'].append('SAFe Agile')
        if 'gitops' in req_lower:
            requirements['methodologies'].append('GitOps')
        if 'devsecops' in req_lower:
            requirements['methodologies'].append('DevSecOps')
        
        # Print what was found
        for category, items in requirements.items():
            if items:
                print(f'  • {category.replace("_", " ").title()}: {", ".join(items)}')
        
        return requirements
    
    def generate_summary_bullets(self, requirements):
        """Generate new summary bullets based on requirements"""
        bullets = []
        
        # Cloud-native architecture
        if requirements['cloud_services']:
            cloud_services = ', '.join(requirements['cloud_services'][:4])
            bullet = f'•   Experience with AWS cloud-native services including {cloud_services} for building scalable microservices architectures and event-driven systems.'
            bullets.append(bullet)
        
        # Microservices & APIs
        if 'microservices' in requirements['other_skills'] or 'BFF' in requirements['other_skills']:
            patterns = []
            if 'BFF' in requirements['other_skills']:
                patterns.append('BFF (Backend for Frontend) microservices')
            if 'microservices' in requirements['other_skills']:
                patterns.append('RESTful APIs')
            
            messaging = ' and '.join([msg for msg in requirements['messaging']][:2])
            if messaging:
                bullet = f'•   Developed {" and ".join(patterns)} for mobile and web applications, implementing event-driven architectures using {messaging} for real-time data integration.'
            else:
                bullet = f'•   Developed {" and ".join(patterns)} for mobile and web applications with focus on scalability and performance optimization.'
            bullets.append(bullet)
        
        # CI/CD
        if requirements['cicd_tools']:
            tools = ', '.join(requirements['cicd_tools'][:3])
            bullet = f'•   Experience in implementing CI/CD pipelines using {tools} for automated software delivery across multiple environments.'
            bullets.append(bullet)
        
        # Methodologies
        if requirements['methodologies']:
            methods = ' and '.join(requirements['methodologies'])
            bullet = f'•   Working experience in {methods} framework environments with cross-functional teams, implementing iterative development practices and continuous improvement processes.'
            bullets.append(bullet)
        
        return bullets
    
    def generate_job_bullets(self, requirements, company_name=''):
        """Generate job-specific experience bullets"""
        bullets = []
        
        # Cloud architecture
        if 'ECS Fargate' in requirements['cloud_services'] or 'Lambda' in requirements['cloud_services']:
            services = []
            if 'ECS Fargate' in requirements['cloud_services']:
                services.append('ECS Fargate')
            if 'Lambda' in requirements['cloud_services']:
                services.append('Lambda')
            
            databases = []
            if 'Aurora PostgreSQL' in requirements['cloud_services']:
                databases.append('Aurora PostgreSQL')
            if 'DynamoDB' in requirements['cloud_services']:
                databases.append('DynamoDB')
            
            if services and databases:
                bullet = f'•   Architected cloud-native applications using AWS {" and ".join(services)} for serverless computing, integrating with {" and ".join(databases)} for scalable data persistence.'
                bullets.append(bullet)
        
        # Event-driven systems
        if requirements['messaging']:
            msg_tools = ' and '.join(requirements['messaging'][:2])
            bullet = f'•   Implemented event-driven microservices using {msg_tools} for real-time data streaming and distributed messaging, processing high-volume events with low latency.'
            bullets.append(bullet)
        
        # BFF pattern
        if 'BFF' in requirements['other_skills']:
            bullet = '•   Designed and deployed BFF (Backend for Frontend) pattern microservices for mobile and web applications, reducing API response latency and improving client-specific data aggregation.'
            bullets.append(bullet)
        
        # CI/CD pipelines
        if 'AWS CodePipeline' in requirements['cloud_services'] or requirements['cicd_tools']:
            tools = []
            if 'AWS CodePipeline' in requirements['cloud_services']:
                tools.append('AWS CodePipeline')
            tools.extend(requirements['cicd_tools'][:2])
            
            if tools:
                bullet = f'•   Created {" and ".join(tools[:2])} workflows for automated CI/CD, implementing blue-green deployments and automated rollback mechanisms.'
                bullets.append(bullet)
        
        # SAFe Agile
        if 'SAFe Agile' in requirements['methodologies']:
            bullet = '•   Established SAFe Agile practices across DevOps teams, conducting PI Planning sessions and implementing continuous improvement through retrospectives and metrics-driven development.'
            bullets.append(bullet)
        
        return bullets[:5]  # Limit to 5 bullets per job
    
    def make_selective_bold(self, paragraph, tech_list):
        """Rebuild paragraph with only tech terms in bold"""
        full_text = paragraph.text
        
        # Clear the paragraph
        for run in paragraph.runs:
            run.text = ''
        
        # Sort tech terms by length (longest first) to match longer terms first
        sorted_tech = sorted(tech_list, key=len, reverse=True)
        
        remaining_text = full_text
        
        while remaining_text:
            # Find the earliest occurrence of any tech term
            earliest_pos = len(remaining_text)
            earliest_term = None
            
            for term in sorted_tech:
                pos = remaining_text.find(term)
                if pos != -1 and pos < earliest_pos:
                    earliest_pos = pos
                    earliest_term = term
            
            if earliest_term:
                # Add text before the term (not bold)
                if earliest_pos > 0:
                    before_text = remaining_text[:earliest_pos]
                    run = paragraph.add_run(before_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.bold = False
                
                # Add the tech term (bold)
                tech_run = paragraph.add_run(earliest_term)
                tech_run.font.name = 'Times New Roman'
                tech_run.font.size = Pt(11)
                tech_run.font.bold = True
                
                # Update remaining text
                remaining_text = remaining_text[earliest_pos + len(earliest_term):]
            else:
                # No more tech terms, add remaining text (not bold)
                if remaining_text:
                    run = paragraph.add_run(remaining_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.bold = False
                break
    
    def insert_summary_bullets(self, bullets):
        """Insert bullets into the Background Summary section"""
        print(f'\n✏️  Adding {len(bullets)} bullets to Summary section...')
        
        # Find insertion point - after Conan/C++ section
        insertion_index = None
        for i, para in enumerate(self.doc.paragraphs):
            if 'Implemented reproducible build workflows by integrating' in para.text and 'Conan' in para.text:
                insertion_index = i + 1
                break
        
        if insertion_index:
            reference_para = self.doc.paragraphs[insertion_index]
            
            for bullet_text in reversed(bullets):
                new_para = reference_para.insert_paragraph_before()
                new_para.text = bullet_text
                
                # Copy formatting
                new_para.paragraph_format.left_indent = reference_para.paragraph_format.left_indent
                new_para.paragraph_format.first_line_indent = reference_para.paragraph_format.first_line_indent
                new_para.paragraph_format.space_before = reference_para.paragraph_format.space_before
                new_para.paragraph_format.space_after = reference_para.paragraph_format.space_after
                
                # Apply selective bold formatting
                self.make_selective_bold(new_para, self.tech_terms)
            
            print('  ✓ Summary section updated')
            return True
        else:
            print('  ✗ Could not find insertion point in summary')
            return False
    
    def insert_job_bullets(self, bullets, company_keyword, year):
        """Insert bullets into a specific job section"""
        print(f'\n✏️  Adding {len(bullets)} bullets to {company_keyword} section...')
        
        # Find the company section
        for i, para in enumerate(self.doc.paragraphs):
            if company_keyword in para.text and year in para.text:
                # Find a good insertion point within the job section
                for j in range(i, min(i+60, len(self.doc.paragraphs))):
                    para_text = self.doc.paragraphs[j].text
                    
                    # Insert after certain key bullets
                    if ('Tekton pipelines with ArgoCD' in para_text or
                        'Integrated Tekton pipelines' in para_text or
                        'Kubernetes for the runtime environment' in para_text or
                        'Flux for GitOps-based cluster state' in para_text):
                        
                        insert_at = j + 1
                        ref_para = self.doc.paragraphs[insert_at]
                        
                        for addition in reversed(bullets):
                            new_p = ref_para.insert_paragraph_before()
                            new_p.text = addition
                            
                            # Copy formatting
                            new_p.paragraph_format.left_indent = ref_para.paragraph_format.left_indent
                            new_p.paragraph_format.first_line_indent = ref_para.paragraph_format.first_line_indent
                            new_p.paragraph_format.space_before = ref_para.paragraph_format.space_before
                            new_p.paragraph_format.space_after = ref_para.paragraph_format.space_after
                            
                            # Apply selective bold formatting
                            self.make_selective_bold(new_p, self.tech_terms)
                        
                        print(f'  ✓ {company_keyword} section updated')
                        return True
                break
        
        print(f'  ✗ Could not find {company_keyword} section')
        return False
    
    def update_technical_skills(self, requirements):
        """Update the Technical Skills table"""
        print('\n✏️  Updating Technical Skills section...')
        
        updates_made = 0
        
        for table in self.doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    category = cells[0].text.strip()
                    content = cells[1].text.strip()
                    
                    # Update based on requirements
                    if 'Cloud Technologies' in category and requirements['cloud_services']:
                        # Add AWS services if not present
                        new_services = [svc for svc in requirements['cloud_services'] if svc not in content]
                        if new_services:
                            # Parse existing content
                            if 'Amazon Web Services' in content:
                                # Add services in parentheses
                                base = 'Amazon Web Services'
                                existing_services = []
                                
                                # Extract existing services if any
                                if '(' in content:
                                    start = content.find('(')
                                    end = content.find(')', start)
                                    if end > start:
                                        existing_services = [s.strip() for s in content[start+1:end].split(',')]
                                
                                # Combine with new services
                                all_services = list(set(existing_services + new_services))
                                services_str = ', '.join(all_services)
                                
                                # Rebuild the cell text
                                cells[1].text = f'Amazon Web Services ({services_str}), Azure, Google Cloud Platform'
                                
                                for para in cells[1].paragraphs:
                                    for run in para.runs:
                                        run.font.name = 'Times New Roman'
                                        run.font.size = Pt(11)
                                
                                updates_made += 1
                    
                    elif 'CI/CD Tools' in category and requirements['cicd_tools']:
                        new_tools = [tool for tool in requirements['cicd_tools'] if tool not in content]
                        if new_tools:
                            cells[1].text = content + ', ' + ', '.join(new_tools)
                            for para in cells[1].paragraphs:
                                for run in para.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                            updates_made += 1
                    
                    elif 'Databases' in category and requirements['databases']:
                        new_dbs = [db for db in requirements['databases'] if db not in content]
                        if new_dbs:
                            cells[1].text = content + ', ' + ', '.join(new_dbs)
                            for para in cells[1].paragraphs:
                                for run in para.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                            updates_made += 1
                    
                    elif 'Web Services' in category and 'BFF' in requirements['other_skills']:
                        if 'BFF' not in content:
                            cells[1].text = content + ', BFF (Backend for Frontend)'
                            for para in cells[1].paragraphs:
                                for run in para.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                            updates_made += 1
            
            # Add Methodologies row if needed
            if requirements['methodologies']:
                has_methodologies = False
                for row in table.rows:
                    if 'Methodologies' in row.cells[0].text:
                        has_methodologies = True
                        break
                
                if not has_methodologies:
                    new_row = table.add_row()
                    new_row.cells[0].text = 'Methodologies'
                    new_row.cells[1].text = ', '.join(requirements['methodologies'])
                    
                    for cell in new_row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                    
                    updates_made += 1
        
        print(f'  ✓ Technical Skills updated ({updates_made} changes)')
    
    def save_resume(self, output_path=None):
        """Save the updated resume"""
        if output_path is None:
            # Create output filename with timestamp in the output directory
            base_name = os.path.splitext(os.path.basename(self.original_resume_path))[0]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'{base_name}_Updated_{timestamp}.docx'
            output_path = os.path.join(self.output_dir, filename)
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
        
        self.doc.save(output_path)
        print(f'\n✅ Updated resume saved: {os.path.basename(output_path)}')
        return output_path
    
    def update_resume(self, requirements_text):
        """Main method to update resume based on requirements"""
        print('\n' + '='*60)
        print('RESUME UPDATER')
        print('='*60)
        
        # Load resume
        self.load_resume()
        
        # Parse requirements
        requirements = self.parse_requirements(requirements_text)
        
        if not any(requirements.values()):
            print('\n⚠️  No relevant technologies found in requirements.')
            print('Please check your input and try again.')
            return None
        
        # Generate new content
        summary_bullets = self.generate_summary_bullets(requirements)
        job_bullets = self.generate_job_bullets(requirements)
        
        # Insert content
        self.insert_summary_bullets(summary_bullets)
        
        # Update most recent job (Early Warning)
        if job_bullets:
            self.insert_job_bullets(job_bullets, 'Early Warning', '2024')
        
        # Update technical skills
        self.update_technical_skills(requirements)
        
        # Save
        output_path = self.save_resume()
        
        print('\n' + '='*60)
        print('UPDATE COMPLETE!')
        print('='*60)
        
        return output_path


def main():
    """Main function"""
    print('='*60)
    print('RESUME UPDATER - DevOps Edition')
    print('='*60)
    print()
    
    # Get the script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Set output directory
    output_dir = '/Users/gokul/Desktop/Devops 12+/edited resumes'
    print(f'Output directory: {output_dir}')
    print()
    
    # Look for job_requirement.txt file
    requirements_file = os.path.join(script_dir, 'job_requirement.txt')
    
    if not os.path.exists(requirements_file):
        print('❌ job_requirement.txt not found!')
        print()
        print('Please create a file named "job_requirement.txt" in the same folder')
        print('and paste your job requirements into it.')
        print()
        print(f'Location: {script_dir}')
        sys.exit(1)
    
    print(f'✓ Found job_requirement.txt')
    
    # Read requirements from file
    with open(requirements_file, 'r', encoding='utf-8') as f:
        requirements_text = f.read()
    
    if not requirements_text.strip():
        print('❌ job_requirement.txt is empty!')
        print('Please paste your job requirements into the file.')
        sys.exit(1)
    
    print(f'✓ Loaded requirements ({len(requirements_text)} characters)')
    
    # Look for resume in same directory
    resume_files = [f for f in os.listdir(script_dir) if f.endswith('.docx') and 'Updated' not in f and '~$' not in f]
    
    if not resume_files:
        print('❌ No resume file found in the current directory.')
        print('Please place your resume (.docx) in the same folder as this script.')
        sys.exit(1)
    
    if len(resume_files) == 1:
        resume_file = resume_files[0]
        print(f'✓ Found resume: {resume_file}')
    else:
        print('Multiple resume files found:')
        for i, f in enumerate(resume_files, 1):
            print(f'  {i}. {f}')
        
        choice = input('\nSelect resume number: ')
        try:
            resume_file = resume_files[int(choice) - 1]
        except (ValueError, IndexError):
            print('Invalid selection.')
            sys.exit(1)
    
    resume_path = os.path.join(script_dir, resume_file)
    
    # Create updater and process
    try:
        updater = ResumeUpdater(resume_path, output_dir)
        output_path = updater.update_resume(requirements_text)
        
        if output_path:
            print(f'\n📄 Updated resume saved to:')
            print(f'   {output_path}')
            print()
            print('Next steps:')
            print('  1. Open the updated resume and review')
            print('  2. Update job_requirement.txt with new requirements')
            print('  3. Run this script again for another job application')
        
    except Exception as e:
        print(f'\n❌ Error: {str(e)}')
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
